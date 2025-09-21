from docx import Document
from docx.shared import Inches

def save_docx(student_name, report_text, chart_path):
    doc = Document()
    # Add new logo to top center
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    run.add_picture('assets/devpro GRT Logo.png', width=Inches(2.5))
    header_paragraph.alignment = 1  # Center

    # Add header text
    doc.add_paragraph("IGNITE 2025 - STUDENT REPORT", style='Title').alignment = 1

    # Add name field
    doc.add_paragraph(f"Name of the Student: {student_name}", style='Heading 2')

    # Add report text
    doc.add_paragraph(report_text)

    # Add chart
    doc.add_picture(chart_path, width=Inches(5))

    filename = f"{student_name}_Report.docx"
    doc.save(filename)
    return filename