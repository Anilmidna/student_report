from openai import OpenAI
from dotenv import load_dotenv
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import os
from concurrent.futures import ThreadPoolExecutor

load_dotenv()  # Load environment variables from .env file
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url="https://api.openai.com/v1"  # explicitly set the base URL for OpenAI platform
)

def make_prompt(student_row):
    name = student_row['Name']
    scores = student_row.drop('Name').to_dict()
    scores_text = ', '.join([f"{subject}: {score}" for subject, score in scores.items()])
    prompt = f"""
Student Name: {name}
Scores: {scores_text}

Write an analysis covering:
- Overall performance
- Strongest subjects
- Weakest subjects
- Improvement suggestions

Make the report detailed and constructive.
"""
    return prompt

def get_report_text(prompt):
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=500
    )
    return response.choices[0].message.content

def generate_chart(student_name, scores):
    plt.figure(figsize=(6,4))
    subjects = list(scores.keys())
    values = list(scores.values())
    plt.bar(subjects, values, color='skyblue')
    plt.title(f"{student_name} - Performance")
    plt.ylabel("Score")
    plt.tight_layout()
    chart_path = os.path.join('reports', f"{student_name}_chart.png")
    os.makedirs('reports', exist_ok=True)
    plt.savefig(chart_path)
    plt.close()
    return chart_path

def save_docx(student_name, report_text, chart_path):
    doc = Document()
    doc.add_heading(f"{student_name}'s Performance Report", 0)
    doc.add_paragraph(report_text)
    doc.add_picture(chart_path, width=Inches(5))
    filename = f"{student_name}_Report.docx"
    doc.save(filename)
    return filename

def docx_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    os.makedirs('reports', exist_ok=True)
    convert(docx_path, pdf_path)
    return pdf_path

def batch_generate_reports(df):
    os.makedirs('reports', exist_ok=True)
    docx_files = []
    with ThreadPoolExecutor(max_workers=5) as executor:
        prompts = [make_prompt(row) for _, row in df.iterrows()]
        futures = [executor.submit(get_report_text, prompt) for prompt in prompts]
        report_texts = [future.result() for future in futures]
    for idx, row in df.iterrows():
        student_name = row['Name']
        scores = row.drop('Name').to_dict()
        chart_path = generate_chart(student_name, scores)

        # Create Word document
        doc = Document()

        # Add new logo to top center in header
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(os.path.join('assets', 'devpro GRT Logo.png'), width=Inches(2.5))
        header_paragraph.alignment = 1  # Center

        # Add header text
        doc.add_paragraph("IGNITE 2025 - STUDENT REPORT", style='Title').alignment = 1

        # Add name field
        doc.add_paragraph(f"Name of the Student: {student_name}", style='Heading 2')

        # Add sections with formatted headings
        doc.add_heading('Performance Analysis', level=1)
        doc.add_paragraph(report_texts[idx])

        # Add chart section
        doc.add_heading('Performance Chart', level=1)
        doc.add_picture(chart_path, width=Inches(6))

        # Add footer
        section = doc.sections[0]
        footer = section.footer
        footer_text = footer.add_paragraph()
        footer_text.text = f"Generated on {pd.Timestamp.now().strftime('%Y-%m-%d')}"
        footer_text.style = doc.styles['Footer']

        # Save document in reports folder
        docx_path = os.path.join('reports', f"{student_name}_report.docx")
        doc.save(docx_path)
        docx_files.append(docx_path)

        # Clean up chart file
        os.remove(chart_path)
    return docx_files